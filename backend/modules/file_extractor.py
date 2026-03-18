"""
modules/file_extractor.py — Universal file-to-text extractor.

Supports: PDF · DOCX · DOC · TXT

The KEY fix in this version: hyperlinks stored as URI annotations in PDFs
and relationship hyperlinks in DOCX are extracted and appended to the text
under a clearly labelled "HYPERLINKS FOUND" section.  The LLM parser then
reads that section to populate github_url, linkedin_url, and other profile
links — even when those links are invisible click-targets on words like
"GitHub" or "Portfolio" in the original document.

Strategy per format
--------------------
PDF  → PyMuPDF (fitz):
         1. page.get_text("text")   — visible text, page by page
         2. page.get_links()        — URI annotations (hyperlinks)
         3. page.get_text("dict")   — fallback scan of any embedded URIs
            in the raw content stream

DOCX → python-docx:
         1. paragraphs + table cells — visible text
         2. XML relationship scan    — all hyperlinks stored in .rels files

DOC  → python-docx fallback, then raw text.
TXT  → UTF-8 / Latin-1 / CP1252 decode chain.
"""

from __future__ import annotations

import io
import logging
import re
from urllib.parse import urlparse

logger = logging.getLogger(__name__)

# ── Domains we specifically care about for profile matching ───────────────
_PROFILE_DOMAINS = (
    "github.com",
    "linkedin.com",
    "leetcode.com",
    "hackerrank.com",
    "codechef.com",
    "codeforces.com",
    "kaggle.com",
    "stackoverflow.com",
    "medium.com",
    "dev.to",
    "behance.net",
    "dribbble.com",
    "portfolio",
    "twitter.com",
    "x.com",
    "gitlab.com",
    "bitbucket.org",
)


def _is_meaningful_url(url: str) -> bool:
    """Return True if the URL is a real external link worth keeping."""
    if not url:
        return False
    url = url.strip()
    # Skip internal PDF anchors and mailto
    if url.startswith(("#", "mailto:", "tel:")):
        return False
    try:
        parsed = urlparse(url)
        return parsed.scheme in ("http", "https") and bool(parsed.netloc)
    except Exception:
        return False


def _label_url(url: str) -> str:
    """Give a human-readable label for a URL based on its domain."""
    url_lower = url.lower()
    if "github.com" in url_lower:
        return f"GitHub: {url}"
    if "linkedin.com" in url_lower:
        return f"LinkedIn: {url}"
    if "leetcode.com" in url_lower:
        return f"LeetCode: {url}"
    if "hackerrank.com" in url_lower:
        return f"HackerRank: {url}"
    if "kaggle.com" in url_lower:
        return f"Kaggle: {url}"
    if "stackoverflow.com" in url_lower:
        return f"StackOverflow: {url}"
    if "medium.com" in url_lower:
        return f"Medium: {url}"
    if "gitlab.com" in url_lower:
        return f"GitLab: {url}"
    if "codechef.com" in url_lower:
        return f"CodeChef: {url}"
    if "codeforces.com" in url_lower:
        return f"Codeforces: {url}"
    if "twitter.com" in url_lower or "x.com" in url_lower:
        return f"Twitter/X: {url}"
    if "behance.net" in url_lower or "dribbble.com" in url_lower:
        return f"Portfolio: {url}"
    return f"Link: {url}"


def _format_links_section(urls: list[str]) -> str:
    """Format a deduplicated list of URLs into an appended section."""
    seen: set[str] = set()
    lines: list[str] = []
    for url in urls:
        url = url.strip().rstrip("/")
        if url and url not in seen and _is_meaningful_url(url):
            seen.add(url)
            lines.append(_label_url(url))
    if not lines:
        return ""
    return "\n\n--- HYPERLINKS FOUND IN DOCUMENT ---\n" + "\n".join(lines)


# ──────────────────────────────────────────────
# PDF
# ──────────────────────────────────────────────

def _extract_pdf(data: bytes) -> str:
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(stream=data, filetype="pdf")
        page_texts: list[str] = []
        all_urls: list[str] = []

        for page in doc:
            # ── Visible text ──────────────────────────────────────────────
            page_texts.append(page.get_text("text"))

            # ── URI annotations (hyperlinks) ──────────────────────────────
            for link in page.get_links():
                uri = link.get("uri", "")
                if uri:
                    all_urls.append(uri)

            # ── Scan raw content stream for any http/https URLs ───────────
            # Catches URLs embedded directly in PDF content that aren't
            # registered as formal link annotations.
            raw = page.get_text("rawdict")
            for block in raw.get("blocks", []):
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        text_val = span.get("text", "")
                        found = re.findall(r'https?://[^\s\)\]>,"\']+', text_val)
                        all_urls.extend(found)

        text = "\n".join(page_texts).strip()

        # Also scan the plain text itself for any written-out URLs
        # (e.g. "github.com/username" without https://)
        plain_urls = re.findall(r'https?://[^\s\)\]>,"\']+', text)
        all_urls.extend(plain_urls)

        # Catch bare URLs like "github.com/user" or "linkedin.com/in/user"
        bare = re.findall(
            r'(?:^|\s)((?:github|linkedin|leetcode|kaggle|hackerrank|'
            r'stackoverflow|medium|gitlab|codechef|codeforces|bitbucket)'
            r'\.(?:com|org|net|io)/[^\s\)\]>,"\']+)',
            text,
            re.IGNORECASE | re.MULTILINE,
        )
        all_urls.extend(["https://" + u.strip() for u in bare])

        if not text:
            logger.warning("PyMuPDF returned empty text — PDF may be image-only.")

        links_section = _format_links_section(all_urls)
        logger.info(
            "PDF extraction: %d chars text, %d unique links found",
            len(text),
            links_section.count("\n") - 1 if links_section else 0,
        )
        return text + links_section

    except ImportError:
        raise RuntimeError("PyMuPDF is not installed. Run: pip install PyMuPDF")
    except Exception as exc:
        raise RuntimeError(f"PDF extraction failed: {exc}") from exc


# ──────────────────────────────────────────────
# DOCX
# ──────────────────────────────────────────────

def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        import zipfile
        from xml.etree import ElementTree as ET

        doc = Document(io.BytesIO(data))
        parts: list[str] = []
        all_urls: list[str] = []

        # ── Visible text — paragraphs ──────────────────────────────────────
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())

        # ── Visible text — tables ──────────────────────────────────────────
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)

        # ── Hyperlinks from DOCX relationships ────────────────────────────
        # DOCX stores hyperlinks in the .rels XML files inside the zip.
        # python-docx exposes them via part.rels.
        try:
            for rel in doc.part.rels.values():
                if "hyperlink" in rel.reltype.lower():
                    target = rel._target
                    if isinstance(target, str):
                        all_urls.append(target)
        except Exception as rel_exc:
            logger.debug("Relationship hyperlink scan failed: %s", rel_exc)

        # ── Direct XML scan for hyperlinks (belt and suspenders) ──────────
        try:
            zf = zipfile.ZipFile(io.BytesIO(data))
            ns = {"r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships"}
            for name in zf.namelist():
                if name.endswith(".rels"):
                    xml = zf.read(name)
                    root = ET.fromstring(xml)
                    for rel in root.findall("Relationship", {}):
                        rtype = rel.get("Type", "")
                        target = rel.get("Target", "")
                        if "hyperlink" in rtype.lower() and target:
                            all_urls.append(target)
        except Exception as zip_exc:
            logger.debug("ZIP hyperlink scan failed: %s", zip_exc)

        text = "\n".join(parts).strip()

        # Also scan visible text for written-out URLs
        plain_urls = re.findall(r'https?://[^\s\)\]>,"\']+', text)
        all_urls.extend(plain_urls)

        links_section = _format_links_section(all_urls)
        return text + links_section

    except ImportError:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx")
    except Exception as exc:
        raise RuntimeError(f"DOCX extraction failed: {exc}") from exc


# ──────────────────────────────────────────────
# TXT
# ──────────────────────────────────────────────

def _extract_txt(data: bytes) -> str:
    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            text = data.decode(encoding).strip()
            # Scan plain text for URLs too
            all_urls = re.findall(r'https?://[^\s\)\]>,"\']+', text)
            links_section = _format_links_section(all_urls)
            return text + links_section
        except UnicodeDecodeError:
            continue
    raise RuntimeError("Could not decode text file with any supported encoding.")


# ──────────────────────────────────────────────
# Public API
# ──────────────────────────────────────────────

def extract_text(filename: str, data: bytes) -> str:
    """
    Extract plain text (plus a hyperlinks section) from an uploaded file.

    Parameters
    ----------
    filename : str
        Original filename including extension (used for format detection).
    data : bytes
        Raw file bytes.

    Returns
    -------
    str
        Extracted plain text with an appended '--- HYPERLINKS FOUND ---'
        section listing all URLs discovered in the document.

    Raises
    ------
    ValueError
        If the file extension is not supported.
    RuntimeError
        If extraction fails (missing library, corrupted file, etc.).
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    logger.info(
        "Extracting text from file: %s (ext=%s, size=%d bytes)",
        filename, ext, len(data),
    )

    if ext == "pdf":
        text = _extract_pdf(data)
    elif ext == "docx":
        text = _extract_docx(data)
    elif ext == "doc":
        try:
            text = _extract_docx(data)
        except Exception:
            logger.warning(".doc extraction via python-docx failed; falling back to raw text.")
            text = _extract_txt(data)
    elif ext == "txt":
        text = _extract_txt(data)
    else:
        raise ValueError(
            f"Unsupported file type '.{ext}'. "
            "Allowed: .pdf · .docx · .doc · .txt"
        )

    if not text:
        raise RuntimeError(
            f"No text could be extracted from '{filename}'. "
            "The file may be empty, image-only, or password-protected."
        )

    logger.info("Extracted %d characters from %s", len(text), filename)
    return text